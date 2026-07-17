import io
import unittest
import zipfile

import pandas as pd
from docx import Document as DocxDocument

from medical_ocr.reporting import (
    block_key,
    build_docx_report,
    build_tables_workbook_xlsx,
    build_tables_zip_csv,
    build_translation_ready_docx,
    dataframe_to_csv_bytes,
    dataframe_to_xlsx_bytes,
    document_is_classified,
    table_block_to_dataframe,
)
from medical_ocr.schema import Block, BlockCategory, BlockType, Document, Page, SourceEngine


def _paragraph_block(text, category=None):
    return Block(
        block_type=BlockType.PARAGRAPH,
        text=text,
        source_engine=SourceEngine.PYMUPDF,
        category=category,
    )


def _table_block(rows, category=None):
    return Block(
        block_type=BlockType.TABLE,
        rows=rows,
        source_engine=SourceEngine.PDFPLUMBER,
        category=category,
    )


class TestTableBlockToDataFrame(unittest.TestCase):
    def test_uses_first_row_as_header_when_no_structured_rows(self):
        block = _table_block([["الدواء", "الجرعة"], ["باراسيتامول", "500 ملغ"]])
        df = table_block_to_dataframe(block)
        self.assertEqual(list(df.columns), ["الدواء", "الجرعة"])
        self.assertEqual(df.iloc[0]["الدواء"], "باراسيتامول")

    def test_uses_structured_rows_when_it_is_a_list(self):
        block = _table_block([["A", "B"], ["1", "2"]])
        structured = [{"الفحص": "Hemoglobin", "القيمة": "13.5"}]
        df = table_block_to_dataframe(block, structured_rows=structured)
        self.assertEqual(list(df.columns), ["الفحص", "القيمة"])
        self.assertEqual(df.iloc[0]["الفحص"], "Hemoglobin")

    def test_falls_back_to_raw_rows_when_structured_rows_is_unparsed_string(self):
        block = _table_block([["A", "B"], ["1", "2"]])
        df = table_block_to_dataframe(block, structured_rows="not json, parse failed")
        self.assertEqual(list(df.columns), ["A", "B"])

    def test_ragged_rows_do_not_crash(self):
        block = _table_block([["A", "B", "C"], ["1", "2"], ["x", "y", "z", "extra"]])
        df = table_block_to_dataframe(block)
        self.assertEqual(len(df), 2)
        self.assertEqual(list(df.columns), ["A", "B", "C"])

    def test_empty_rows_returns_empty_dataframe_without_indexerror(self):
        block = Block(block_type=BlockType.TABLE, rows=[], source_engine=SourceEngine.PDFPLUMBER)
        df = table_block_to_dataframe(block)
        self.assertTrue(df.empty)


class TestDataframeExport(unittest.TestCase):
    def test_csv_bytes_start_with_utf8_bom(self):
        df = pd.DataFrame([{"اسم": "أحمد"}])
        self.assertEqual(dataframe_to_csv_bytes(df)[:3], b"\xef\xbb\xbf")

    def test_csv_bytes_contain_arabic_text_correctly_decoded(self):
        df = pd.DataFrame([{"اسم": "أحمد محمد"}])
        decoded = dataframe_to_csv_bytes(df).decode("utf-8-sig")
        self.assertIn("أحمد محمد", decoded)

    def test_xlsx_bytes_roundtrip_via_openpyxl(self):
        df = pd.DataFrame([{"الفحص": "Hemoglobin", "القيمة": "13.5"}])
        xlsx_bytes = dataframe_to_xlsx_bytes(df)
        roundtripped = pd.read_excel(io.BytesIO(xlsx_bytes))
        self.assertEqual(roundtripped.iloc[0]["الفحص"], "Hemoglobin")
        self.assertEqual(str(roundtripped.iloc[0]["القيمة"]), "13.5")


def _classified_document():
    page1 = Page(
        page_number=1,
        source="digital",
        blocks=[
            _paragraph_block("اسم المريض: أحمد", category=BlockCategory.PATIENT_INFO),
            _table_block(
                [["الفحص", "القيمة"], ["Hemoglobin", "13.5"]],
                category=BlockCategory.CLINICAL_RESULTS,
            ),
        ],
    )
    page2 = Page(
        page_number=2,
        source="digital",
        blocks=[
            _paragraph_block("يُنصح بمتابعة الحالة أسبوعياً", category=BlockCategory.DOCTOR_NOTES),
        ],
    )
    return Document(file_name="test.pdf", pages=[page1, page2])


class TestBuildDocxReport(unittest.TestCase):
    def test_groups_by_category_when_classified(self):
        docx_bytes = build_docx_report(_classified_document())
        reopened = DocxDocument(io.BytesIO(docx_bytes))
        headings = [p.text for p in reopened.paragraphs if p.style.name.startswith("Heading")]
        self.assertIn("معلومات المريض", headings)
        self.assertIn("النتائج السريرية", headings)
        self.assertIn("ملاحظات الطبيب", headings)

    def test_skips_category_with_zero_blocks(self):
        docx_bytes = build_docx_report(_classified_document())
        reopened = DocxDocument(io.BytesIO(docx_bytes))
        headings = [p.text for p in reopened.paragraphs if p.style.name.startswith("Heading")]
        self.assertNotIn("غير مصنّف", headings)

    def test_falls_back_to_page_order_when_nothing_classified(self):
        page = Page(page_number=1, source="digital", blocks=[_paragraph_block("نص عادي")])
        document = Document(file_name="test.pdf", pages=[page])
        self.assertFalse(document_is_classified(document))

        docx_bytes = build_docx_report(document)
        reopened = DocxDocument(io.BytesIO(docx_bytes))
        headings = [p.text for p in reopened.paragraphs if p.style.name.startswith("Heading")]
        self.assertEqual(headings, ["صفحة 1"])

    def test_renders_table_as_real_docx_table_not_flattened_text(self):
        docx_bytes = build_docx_report(_classified_document())
        reopened = DocxDocument(io.BytesIO(docx_bytes))
        self.assertEqual(len(reopened.tables), 1)
        table = reopened.tables[0]
        self.assertEqual(table.rows[0].cells[0].text, "الفحص")
        self.assertEqual(table.rows[1].cells[0].text, "Hemoglobin")

    def test_uses_provided_structured_rows_over_raw_rows(self):
        document = _classified_document()
        key = block_key(1, 1)
        docx_bytes = build_docx_report(
            document, structured_tables={key: [{"الفحص": "Glucose", "القيمة": "95"}]}
        )
        reopened = DocxDocument(io.BytesIO(docx_bytes))
        table = reopened.tables[0]
        self.assertEqual(table.rows[1].cells[0].text, "Glucose")

    def test_falls_back_to_raw_rows_when_structured_tables_value_is_not_a_list(self):
        document = _classified_document()
        key = block_key(1, 1)
        docx_bytes = build_docx_report(document, structured_tables={key: "not a list"})
        reopened = DocxDocument(io.BytesIO(docx_bytes))
        table = reopened.tables[0]
        self.assertEqual(table.rows[1].cells[0].text, "Hemoglobin")

    def test_empty_document_produces_valid_docx_without_crashing(self):
        document = Document(file_name="empty.pdf", pages=[])
        docx_bytes = build_docx_report(document)
        reopened = DocxDocument(io.BytesIO(docx_bytes))
        self.assertEqual(len(reopened.tables), 0)


class TestBuildTranslationReadyDocx(unittest.TestCase):
    def test_raises_when_not_classified(self):
        page = Page(page_number=1, source="digital", blocks=[_paragraph_block("نص")])
        document = Document(file_name="test.pdf", pages=[page])
        with self.assertRaises(ValueError):
            build_translation_ready_docx(document)

    def test_excludes_page_prefix_and_includes_patient_and_doctor_text(self):
        docx_bytes = build_translation_ready_docx(_classified_document())
        reopened = DocxDocument(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in reopened.paragraphs)
        self.assertNotIn("[صفحة", all_text)
        self.assertIn("اسم المريض: أحمد", all_text)
        self.assertIn("يُنصح بمتابعة الحالة أسبوعياً", all_text)

    def test_clinical_results_rendered_as_table(self):
        docx_bytes = build_translation_ready_docx(_classified_document())
        reopened = DocxDocument(io.BytesIO(docx_bytes))
        self.assertEqual(len(reopened.tables), 1)
        self.assertEqual(reopened.tables[0].rows[0].cells[0].text, "الفحص")

    def test_other_category_is_excluded(self):
        page = Page(
            page_number=1,
            source="digital",
            blocks=[
                _paragraph_block("معلومة مريض", category=BlockCategory.PATIENT_INFO),
                _paragraph_block("محتوى غير مصنّف بوضوح", category=BlockCategory.OTHER),
            ],
        )
        document = Document(file_name="test.pdf", pages=[page])
        docx_bytes = build_translation_ready_docx(document)
        reopened = DocxDocument(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in reopened.paragraphs)
        self.assertNotIn("محتوى غير مصنّف بوضوح", all_text)


class TestBuildTablesWorkbookXlsx(unittest.TestCase):
    def test_one_sheet_per_table(self):
        xlsx_bytes = build_tables_workbook_xlsx(_classified_document())
        sheets = pd.read_excel(io.BytesIO(xlsx_bytes), sheet_name=None)
        self.assertEqual(len(sheets), 1)

    def test_raises_when_no_tables(self):
        page = Page(page_number=1, source="digital", blocks=[_paragraph_block("نص فقط")])
        document = Document(file_name="test.pdf", pages=[page])
        with self.assertRaises(ValueError):
            build_tables_workbook_xlsx(document)

    def test_sheet_names_respect_length_limit(self):
        xlsx_bytes = build_tables_workbook_xlsx(_classified_document())
        sheets = pd.read_excel(io.BytesIO(xlsx_bytes), sheet_name=None)
        for sheet_name in sheets:
            self.assertLessEqual(len(sheet_name), 31)


class TestBuildTablesZipCsv(unittest.TestCase):
    def test_one_csv_per_table(self):
        zip_bytes = build_tables_zip_csv(_classified_document())
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as archive:
            names = archive.namelist()
        self.assertEqual(len(names), 1)
        self.assertTrue(names[0].endswith(".csv"))

    def test_raises_when_no_tables(self):
        page = Page(page_number=1, source="digital", blocks=[_paragraph_block("نص فقط")])
        document = Document(file_name="test.pdf", pages=[page])
        with self.assertRaises(ValueError):
            build_tables_zip_csv(document)


if __name__ == "__main__":
    unittest.main()
