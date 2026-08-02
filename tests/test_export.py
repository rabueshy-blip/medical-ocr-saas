"""اختبار ميزة استخراج الأصول (Assets): تصدير Word بلا صور يُعيد .docx مفرداً كما
كان. وجود صور يُعيد ZIP واحداً (Word + مجلد images/) — الصورة الحقيقية مُضمَّنة الآن
في مكانها داخل Word نفسه (قرار مُحدَّث)، ونسخة إضافية بدقة كاملة في مجلد images/ —
راجع `medical_ocr/api/routers/export.py`."""

import io
import unittest
import zipfile

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from medical_ocr.api.app import app


class TestExportDocx(unittest.TestCase):
    def setUp(self):
        self.client = TestClient(app)

    def test_export_without_images_returns_plain_docx(self):
        content = {
            "type": "doc",
            "content": [
                {"type": "paragraph", "content": [{"type": "text", "text": "Hello"}]},
            ],
        }

        response = self.client.post(
            "/export-docx", json={"content": content, "file_name": "no_images", "images": []}
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        docx_doc = DocxDocument(io.BytesIO(response.content))
        self.assertEqual(docx_doc.paragraphs[0].text, "Hello")

    # PNG أحمر 4×3 بكسل حقيقي وصالح — لازم لأن `add_picture` (خلاف الإصدار القديم الذي
    # كان يكتب placeholder نصياً فقط) يمرّر البيانات فعلياً لمحلّل صور Pillow/python-docx.
    _TINY_PNG_BASE64 = (
        "iVBORw0KGgoAAAANSUhEUgAAAAQAAAADCAIAAAA7ljmRAAAAEUlEQVR4"
        "nGP8z4AATEhsVA4AJnYBBZNezToAAAAASUVORK5CYII="
    )

    def test_export_with_images_embeds_real_picture_and_ships_png_folder(self):
        content = {
            "type": "doc",
            "content": [
                {"type": "paragraph", "content": [{"type": "text", "text": "Before"}]},
                {
                    "type": "image",
                    "attrs": {
                        "src": f"data:image/png;base64,{self._TINY_PNG_BASE64}",
                        "imageId": "Image_01",
                    },
                },
                {"type": "paragraph", "content": [{"type": "text", "text": "After"}]},
            ],
        }
        images = [{"image_id": "Image_01", "mime_type": "image/png", "data_base64": self._TINY_PNG_BASE64}]

        response = self.client.post(
            "/export-docx", json={"content": content, "file_name": "with_images", "images": images}
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.headers["content-type"], "application/zip")

        with zipfile.ZipFile(io.BytesIO(response.content)) as zf:
            names = zf.namelist()
            self.assertIn("with_images.docx", names)
            self.assertIn("images/Image_01.png", names)

            with zf.open("with_images.docx") as f:
                docx_doc = DocxDocument(io.BytesIO(f.read()))
                paragraph_texts = [p.text for p in docx_doc.paragraphs]
                # لا نص placeholder — الصورة عنصر InlineShape حقيقي بين الفقرتين، ليست سطر نص.
                self.assertEqual(paragraph_texts, ["Before", "", "After"])
                self.assertEqual(len(docx_doc.inline_shapes), 1)

            with zf.open("images/Image_01.png") as f:
                self.assertTrue(f.read().startswith(b"\x89PNG"))

    def test_export_with_undecodable_image_src_falls_back_to_placeholder(self):
        # src غير بصيغة data URL متوقَّعة (حالة دفاعية غير متوقَّعة عملياً) — يجب أن يكتب
        # نفس عبارة الـplaceholder القديمة بدل فشل تصدير المستند بالكامل.
        content = {
            "type": "doc",
            "content": [
                {"type": "image", "attrs": {"src": "not-a-data-url", "imageId": "Image_02"}},
            ],
        }

        response = self.client.post(
            "/export-docx", json={"content": content, "file_name": "bad_src", "images": []}
        )

        self.assertEqual(response.status_code, 200)
        docx_doc = DocxDocument(io.BytesIO(response.content))
        self.assertEqual(docx_doc.paragraphs[0].text, "[Insert Image_02 here]")
        self.assertEqual(len(docx_doc.inline_shapes), 0)

    def test_arabic_file_name_does_not_crash_response_headers(self):
        # اسم ملف عربي (شائع جداً لملفات المستخدم الفعلية) كان يمر دون تغيير عبر
        # `_safe_file_name` (لأن `\w` في Python يطابق يونيكود أيضاً)، فيُسبِّب
        # UnicodeEncodeError عند وضعه في ترويسة Content-Disposition (يجب أن تكون
        # latin-1) — اختُبِر هذا الخطأ فعلياً عبر الواجهة الحية قبل إصلاحه.
        content = {"type": "doc", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Hi"}]}]}

        response = self.client.post(
            "/export-docx",
            json={"content": content, "file_name": "تقرير_DEXA_للمريضة", "images": []},
        )

        self.assertEqual(response.status_code, 200)
        response.headers["content-disposition"]  # لا يرفع استثناءً عند الوصول إليها


if __name__ == "__main__":
    unittest.main()
