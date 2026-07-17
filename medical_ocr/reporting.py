"""
طبقة تحويل بحتة (لا تستدعي أي LM، ولا تستورد Streamlit) من Document/Block إلى
صيغ تصدير: pandas.DataFrame للجداول، CSV/Excel، وتقارير Word (.docx) — تُستخدَم
من `streamlit_app.py` عبر أزرار "خيارات التنزيل" (القسم الرابع من طلب المستخدم
لواجهة "احترافية").

ثلاثة صيغ Word/Excel مختلفة حسب حالة استخدام مختلفة:
- `build_docx_report`: "التقرير الكامل" — كل المحتوى، مُصنَّفاً إن وُجد تصنيف
  وإلا بترتيب الصفحات (fallback رشيق، لا يفشل أبداً).
- `build_translation_ready_docx`: "نص جاهز للترجمة" — لمترجمين يستخدمون أدوات
  CAT (مثل MateCat) التي تُقسِّم النص إلى شرائح (segments) بحسب الفقرات؛ لذلك
  فقرات معلومات المريض/ملاحظات الطبيب تُكتب دون أي تنسيق إضافي (لا بادئة "[صفحة
  N]"، لا خط عريض) قد يُفسد تقسيم الشرائح، بينما تبقى جداول النتائج السريرية
  جداول Word حقيقية (بيانات مرجعية غير مترجَمة) بدل تحويلها لنص مسطَّح. يتطلب
  تصنيفاً فعلياً (لا يوجد fallback منطقي هنا خلافاً للتقرير الكامل).
- `build_tables_workbook_xlsx` / `build_tables_zip_csv`: "بيانات الجداول" —
  كل الجداول المكتشَفة في مكان واحد (ملف Excel متعدد الأوراق، أو أرشيف zip من
  ملفات CSV) بدل زر تنزيل منفصل لكل جدول على حدة.
"""

from __future__ import annotations

import io
import zipfile
from typing import Dict, List, Optional

import pandas as pd
from docx import Document as DocxDocument

from .schema import Block, BlockCategory, BlockType, Document

CATEGORY_ORDER = [
    BlockCategory.PATIENT_INFO,
    BlockCategory.CLINICAL_RESULTS,
    BlockCategory.DOCTOR_NOTES,
    BlockCategory.OTHER,
]

CATEGORY_LABELS_AR = {
    BlockCategory.PATIENT_INFO: "معلومات المريض",
    BlockCategory.CLINICAL_RESULTS: "النتائج السريرية",
    BlockCategory.DOCTOR_NOTES: "ملاحظات الطبيب",
    BlockCategory.OTHER: "غير مصنّف",
}

_MAX_XLSX_SHEET_NAME_LEN = 31


def block_key(page_number: int, block_index: int) -> str:
    """نفس اتفاقية block_key المستخدَمة في `streamlit_app.py` لقاموس corrections
    (`f"p{page_number}_b{block_index}"}`) — مركزيّة هنا كي لا تتكرر بصيغتين مختلفتين."""
    return f"p{page_number}_b{block_index}"


def document_is_classified(document: Document) -> bool:
    """مصدر الحقيقة الوحيد لسؤال 'هل صُنِّف هذا المستند؟' — تُستخدَم لتحديد ما إذا
    كان التقرير الكامل يُبنى حسب الفئة أو حسب ترتيب الصفحات، وما إذا كانت واجهة
    Streamlit تعرض تبويبات الفئات أو العرض المسطَّح الحالي."""
    return any(block.category is not None for page in document.pages for block in page.blocks)


def table_block_to_dataframe(block: Block, structured_rows: object = None) -> pd.DataFrame:
    """يحوّل Block من نوع جدول إلى DataFrame — يفضّل structured_rows (ناتج AI
    مُهيكَل، list[dict]) إن كان صالحاً، وإلا يعتمد على block.rows الخام (الصف
    الأول عنوان، البقية بيانات)، مع حماية من صفوف غير منتظمة الطول (شائعة في
    جداول OCR/pdfplumber) كي لا يرفع pandas استثناءً على عدم تطابق الطول."""
    if isinstance(structured_rows, list) and structured_rows:
        return pd.DataFrame(structured_rows)

    rows = block.rows or []
    if not rows:
        return pd.DataFrame()

    header, *data_rows = rows
    width = len(header)
    normalized_data = [
        (row + [""] * width)[:width] if len(row) < width else row[:width] for row in data_rows
    ]
    return pd.DataFrame(normalized_data, columns=header)


def dataframe_to_csv_bytes(df: pd.DataFrame) -> bytes:
    # utf-8-sig (BOM) إلزامي: Excel يعرض النص العربي كرموز غير مفهومة (mojibake)
    # بدون BOM عند فتح CSV بترميز utf-8 عادي.
    return df.to_csv(index=False).encode("utf-8-sig")


def dataframe_to_xlsx_bytes(df: pd.DataFrame, sheet_name: str = "الجدول") -> bytes:
    buffer = io.BytesIO()
    df.to_excel(buffer, index=False, sheet_name=sheet_name, engine="openpyxl")
    return buffer.getvalue()


def _resolve_table_rows(block: Block, block_index: int, page_number: int, structured_tables: Optional[Dict[str, object]]):
    structured_tables = structured_tables or {}
    structured = structured_tables.get(block_key(page_number, block_index))
    return structured if isinstance(structured, list) and structured else None


def _iter_all_table_blocks(document: Document):
    """يولّد (page_number, block_index, block) لكل Block من نوع جدول في المستند كاملاً."""
    for page in document.pages:
        for index, block in enumerate(page.blocks):
            if block.block_type == BlockType.TABLE:
                yield page.page_number, index, block


def _add_table_to_docx(docx_document: DocxDocument, df: pd.DataFrame) -> None:
    if df.empty:
        docx_document.add_paragraph("(جدول فارغ)")
        return
    table = docx_document.add_table(rows=1, cols=len(df.columns))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    header_cells = table.rows[0].cells
    for col_index, column_name in enumerate(df.columns):
        header_cells[col_index].text = str(column_name)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            cells[col_index].text = "" if pd.isna(value) else str(value)


def build_docx_report(document: Document, structured_tables: Optional[Dict[str, object]] = None) -> bytes:
    """'التقرير الكامل (Word)' — مُصنَّف حسب الفئة إن وُجد تصنيف، وإلا حسب ترتيب
    الصفحات (fallback رشيق، لا يفشل أبداً حتى لو لم يُشغَّل التصنيف بعد)."""
    docx_document = DocxDocument()

    if document_is_classified(document):
        for category in CATEGORY_ORDER:
            category_blocks = [
                (page.page_number, index, block)
                for page in document.pages
                for index, block in enumerate(page.blocks)
                if block.category == category
            ]
            if not category_blocks:
                continue
            docx_document.add_heading(CATEGORY_LABELS_AR[category], level=1)
            for page_number, index, block in category_blocks:
                _add_block_to_docx(docx_document, block, page_number, index, structured_tables, page_prefix=True)
    else:
        for page in document.pages:
            docx_document.add_heading(f"صفحة {page.page_number}", level=1)
            for index, block in enumerate(page.blocks):
                _add_block_to_docx(docx_document, block, page.page_number, index, structured_tables, page_prefix=False)

    buffer = io.BytesIO()
    docx_document.save(buffer)
    return buffer.getvalue()


def _add_block_to_docx(
    docx_document: DocxDocument,
    block: Block,
    page_number: int,
    block_index: int,
    structured_tables: Optional[Dict[str, object]],
    page_prefix: bool,
) -> None:
    if block.block_type == BlockType.TABLE:
        if page_prefix:
            docx_document.add_paragraph(f"[صفحة {page_number}]").runs[0].bold = True
        structured = _resolve_table_rows(block, block_index, page_number, structured_tables)
        df = table_block_to_dataframe(block, structured)
        _add_table_to_docx(docx_document, df)
    else:
        paragraph = docx_document.add_paragraph()
        if page_prefix:
            prefix_run = paragraph.add_run(f"[صفحة {page_number}] ")
            prefix_run.bold = True
        paragraph.add_run(block.text or "")


def build_translation_ready_docx(document: Document, structured_tables: Optional[Dict[str, object]] = None) -> bytes:
    """'نص جاهز للترجمة (DOCX)' — يتطلب تصنيفاً فعلياً (لا fallback منطقي هنا).

    معلومات المريض/ملاحظات الطبيب: فقرات نظيفة بلا أي تنسيق إضافي (لا بادئة صفحة،
    لا خط عريض) كي لا تُفسِد أدوات CAT (مثل MateCat) تقسيم النص إلى شرائح.
    النتائج السريرية: جداول Word حقيقية (بيانات مرجعية غير مترجَمة)، وليست نصاً
    مسطَّحاً، تحت عنوان واحد في نهاية المستند حتى يبقى الملف صالحاً بنيوياً.
    محتوى 'غير مصنّف' يُستبعَد بالكامل — غير مرتبط بمهمة الترجمة.
    """
    if not document_is_classified(document):
        raise ValueError("لا يمكن توليد نص جاهز للترجمة قبل تشغيل تصنيف المستند.")

    docx_document = DocxDocument()

    for category in (BlockCategory.PATIENT_INFO, BlockCategory.DOCTOR_NOTES):
        for page in document.pages:
            for index, block in enumerate(page.blocks):
                if block.category != category or block.block_type == BlockType.TABLE:
                    continue
                docx_document.add_paragraph(block.text or "")

    clinical_tables = [
        (page.page_number, index, block)
        for page in document.pages
        for index, block in enumerate(page.blocks)
        if block.category == BlockCategory.CLINICAL_RESULTS and block.block_type == BlockType.TABLE
    ]
    if clinical_tables:
        docx_document.add_heading("بيانات سريرية (مرجعية)", level=1)
        for page_number, index, block in clinical_tables:
            structured = _resolve_table_rows(block, index, page_number, structured_tables)
            df = table_block_to_dataframe(block, structured)
            _add_table_to_docx(docx_document, df)

    buffer = io.BytesIO()
    docx_document.save(buffer)
    return buffer.getvalue()


def _unique_sheet_name(base_name: str, used_names: set) -> str:
    name = base_name[:_MAX_XLSX_SHEET_NAME_LEN]
    suffix = 1
    while name in used_names:
        suffix_str = f"_{suffix}"
        name = base_name[: _MAX_XLSX_SHEET_NAME_LEN - len(suffix_str)] + suffix_str
        suffix += 1
    used_names.add(name)
    return name


def build_tables_workbook_xlsx(document: Document, structured_tables: Optional[Dict[str, object]] = None) -> bytes:
    """'بيانات الجداول (Excel)' — ملف واحد متعدد الأوراق، ورقة لكل جدول مكتشَف
    في المستند كاملاً، بدل زر تنزيل منفصل لكل جدول."""
    tables = list(_iter_all_table_blocks(document))
    if not tables:
        raise ValueError("لا توجد جداول في هذا المستند للتصدير.")

    buffer = io.BytesIO()
    used_names: set = set()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for page_number, index, block in tables:
            structured = _resolve_table_rows(block, index, page_number, structured_tables)
            df = table_block_to_dataframe(block, structured)
            sheet_name = _unique_sheet_name(block_key(page_number, index), used_names)
            df.to_excel(writer, index=False, sheet_name=sheet_name)
    return buffer.getvalue()


def build_tables_zip_csv(document: Document, structured_tables: Optional[Dict[str, object]] = None) -> bytes:
    """'بيانات الجداول (CSV)' — أرشيف zip واحد يحوي ملف CSV لكل جدول مكتشَف."""
    tables = list(_iter_all_table_blocks(document))
    if not tables:
        raise ValueError("لا توجد جداول في هذا المستند للتصدير.")

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for page_number, index, block in tables:
            structured = _resolve_table_rows(block, index, page_number, structured_tables)
            df = table_block_to_dataframe(block, structured)
            archive.writestr(f"{block_key(page_number, index)}.csv", dataframe_to_csv_bytes(df))
    return buffer.getvalue()
